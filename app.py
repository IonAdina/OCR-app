from flask import Flask, render_template, request, jsonify
from PIL import Image
import pytesseract
from docx import Document
import PyPDF2

app = Flask(__name__)

pytesseract.pytesseract.tesseract_cmd = r'C:\Users\Adina\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    result = {"text": []}

    for i in range(1, 100):  # Schimbă 6 cu numărul total de fișiere
        file_key = f'file{i}'
        if file_key in request.files:
            file = request.files[file_key]
            file_result = {"filename": file.filename, "content": ""}

            if file.filename.endswith('.docx'):
                # Procesează documentul Word
                text = extract_text_from_docx(file)
                document_type = identify_document_type(text)
                file_result["content"] = text
                file_result["document_type"] = document_type
                
                if document_type == "Anexa nr. 1 pentru bursa sociala":
                    nume_student, facultate, specializare, an, grupa, medie,CodNP_student, AnUniv, motiv,VNetTotal,pensii,alocatii,ajutoare_stat,spatii_inchiriate,nr_persoane,Venit_per_membru, nr_credite = filter_text_for_anexa1(text)
                    file_result["nume_student"] = nume_student
                    file_result["facultate"] = facultate
                    file_result["specializare"] = specializare
                    file_result["an"] = an
                    file_result["grupa"] = grupa
                    file_result["medie"] = medie
                    file_result["CodNP_student"] = CodNP_student
                    file_result["AnUniv"] = AnUniv
                    file_result["motiv"] = motiv
                    file_result["VNetTotal"] = VNetTotal
                    file_result["pensii"] = pensii
                    file_result["alocatii"] = alocatii
                    file_result["ajutoare_stat"]=ajutoare_stat
                    file_result["spatii_inchiriate"]= spatii_inchiriate
                    file_result["nr_persoane"]= nr_persoane
                    file_result["Venit_per_membru"] = Venit_per_membru
                    file_result["nr_credite"]= nr_credite

                if document_type == "Anexa nr. 2 pentru bursa sociala":
                    nume_student, facultate, specializare, an, grupa, medie, CodNP_student, AnUniv ,perioada,nr_credite= filter_text_for_anexa2(text)
                    file_result["nume_student"] = nume_student
                    file_result["facultate"] = facultate
                    file_result["specializare"] = specializare
                    file_result["an"] = an
                    file_result["grupa"] = grupa
                    file_result["medie"] = medie
                    file_result["CodNP_student"] = CodNP_student
                    file_result["AnUniv"] = AnUniv
                    file_result["perioada"] = perioada
                    file_result["nr_credite"]= nr_credite

                if document_type == "Anexa nr. 3 pentru toate tipurile de burse, cu excepția bursei sociale":
                    nume_student, facultate, specializare, an, grupa, medie, CodNP_student, AnUniv,tip_bursa,motiv,nr_credite= filter_text_for_anexa3(text)
                    file_result["nume_student"] = nume_student
                    file_result["facultate"] = facultate
                    file_result["specializare"] = specializare
                    file_result["an"] = an
                    file_result["grupa"] = grupa
                    file_result["medie"] = medie
                    file_result["CodNP_student"] = CodNP_student
                    file_result["AnUniv"] = AnUniv
                    file_result["tip_bursa"] = tip_bursa
                    file_result["motiv"]=motiv
                    file_result["nr_credite"]=nr_credite

                if document_type == "Acord de prelucrarea datelor cu caracter personal al membrilor familiei studentului solicitant de bursă socială":
                    nume_parinte,data_nastere,CodNP_parinte,tip_membru,nume_student,facultate,specializare,tip_bursa,adresa,telefon_parinte= filter_acord_membrii_fam(text)
                    file_result["nume_parinte"] = nume_parinte
                    file_result["data_nastere"] = data_nastere
                    file_result["CodNP_parinte"] = CodNP_parinte
                    file_result["tip_membru"] = tip_membru
                    file_result["nume_student"] = nume_student
                    file_result["facultate"] = facultate
                    file_result["specializare"] = specializare
                    file_result["tip_bursa"] = tip_bursa
                    file_result["adresa"] = adresa
                    file_result["telefon_parinte"]=telefon_parinte

                if document_type =="Acord de prelucrarea datelor cu caracter personal al solicitantului de bursă":
                    nume_student,facultate,mama_student,tata_student,domiciliu_student,telefon_student=filter_acord_solicitant(text)
                    file_result["nume_student"] = nume_student
                    file_result["facultate"] = facultate
                    file_result["mama_student"] = mama_student
                    file_result["tata_student"] = tata_student
                    file_result["nume_student"] = nume_student
                    file_result["domiciliu_student"] = domiciliu_student
                    file_result["telefon_student"] = telefon_student

                if document_type =="Declarație de acceptare a bursei de Performanță":
                    nume_student,facultate,program_studii,AnUniv,cont_deschis_la,titular,iban=filter_text_anexa4(text)
                    file_result["nume_student"] = nume_student
                    file_result["facultate"] = facultate
                    file_result["program_studii"] = program_studii
                    file_result["AnUniv"] = AnUniv
                    file_result["cont_deschis_la"] = cont_deschis_la
                    file_result["titular"] = titular
                    file_result["iban"] = iban



            elif file.filename.endswith(('.jpg', '.jpeg', '.png')):
                # Procesează imaginea utilizând Tesseract
                text = extract_text_from_image(file)
                file_result["content"] = text

            elif file.filename.endswith('.pdf'):
                # Procesează PDF-ul
                text = extract_text_from_pdf(file)
                file_result["content"] = text

            else:
                text = "Formatul fișierului nu este suportat."
                file_result["content"] = text

            result["text"].append(file_result)

    return jsonify(result)

def filter_text_for_anexa1(text):
    start_index1 = text.find("Subsemnatul(a)") + len("Subsemnatul(a)")
    end_index1 = text.find(",student(a)")
    if start_index1 != -1 and end_index1 != -1:
        nume_student = text[start_index1:end_index1].strip()
    
    start_index2 = text.find("Facultății de") + len("Facultății de")
    end_index2 = text.find(",domeniul/specializarea")
    if start_index2 != -1 and end_index2 != -1:
        facultate = text[start_index2:end_index2].strip()

    start_index3 = text.find(",domeniul/specializarea") + len(",domeniul/specializarea")
    end_index3 = text.find(",în")
    if start_index3 != -1 and end_index3 != -1:
        specializare = text[start_index3:end_index3].strip()

    start_index4 = text.find("anul") + len("anul")
    end_index4 = text.find(",grupa")
    if start_index4 != -1 and end_index4 != -1:
        an = text[start_index4:end_index4].strip()

    start_index5 = text.find(",grupa") + len(",grupa")
    end_index5 = text.find(",media")
    if start_index5 != -1 and end_index5 != -1:
        grupa = text[start_index5:end_index5].strip()

    start_index6 = text.find(",media") + len(",media")
    end_index6 = text.find(",numărul")
    if start_index6 != -1 and end_index6 != -1:
        medie = text[start_index6:end_index6].strip()
    

    start_index7 = text.find(",CNP") + len(",CNP")
    end_index7 = text.find(",rog")
    if start_index7 != -1 and end_index7 != -1:
        CodNP_student = text[start_index7:end_index7].strip()

    start_index8 = text.find("universitar") + len("universitar")
    end_index8 = text.find(",a")
    if start_index8 != -1 and end_index8 != -1:
        AnUniv = text[start_index8:end_index8].strip()

    start_index9 = text.find("motive:") + len("motive:")
    end_index9 = text.find("În vederea")
    if start_index9 != -1 and end_index9 != -1:
        motiv = text[start_index9:end_index9].strip()

    start_index10 = text.find("totale") + len("totale")
    end_index10 = text.find("Pensii")
    if start_index10 != -1 and end_index10 != -1:
        VNetTotal = text[start_index10:end_index10].strip()
    VNetTotal=VNetTotal[:len(VNetTotal)-1]

    start_index11 = text.find("Pensii") + len("Pensii")
    end_index11 = text.find("Alocații")
    if start_index11 != -1 and end_index11 != -1:
        pensii = text[start_index11:end_index11].strip()
    pensii=pensii[:len(pensii)-1]

    start_index12 = text.find("pentru copii") + len("pentru copii")
    end_index12 = text.find("Alte ajutoare")
    if start_index12 != -1 and end_index12 != -1:
        alocatii = text[start_index12:end_index12].strip()
    alocatii=alocatii[:len(alocatii)-1]

    start_index13 = text.find("primite de la stat") + len("primite de la stat")
    end_index13 = text.find("Venituri din")
    if start_index13 != -1 and end_index13 != -1:
        ajutoare_stat = text[start_index13:end_index13].strip()
    ajutoare_stat=ajutoare_stat[:len(ajutoare_stat)-1]

    start_index14 = text.find("spații închiriate") + len("spații închiriate")
    end_index14 = text.find("B.")
    if start_index14 != -1 and end_index14 != -1:
        spatii_inchiriate = text[start_index14:end_index14].strip()
    spatii_inchiriate=spatii_inchiriate[:len(spatii_inchiriate)-1]

    start_index15 = text.find("aflate în întreținere") + len("aflate în întreținere")
    end_index15= text.find("Numărul elevilor")
    if start_index15 != -1 and end_index15 != -1:
        nr_persoane = text[start_index15:end_index15].strip()

    start_index16 = text.find("membru de familie (C=A/B)") + len("membru de familie (C=A/B)")
    end_index16 = text.find("Subsemnatul _____________________,")
    if start_index16 != -1 and end_index16 != -1:
        Venit_per_membru = text[start_index16:end_index16].strip()
    
    start_index17 = text.find("de credite") + len("de credite")
    end_index17 = text.find(",CNP")
    if start_index17 != -1 and end_index17 != -1:
        nr_credite = text[start_index17:end_index17].strip()
    
    return nume_student, facultate, specializare, an ,grupa , medie, CodNP_student, AnUniv, motiv, VNetTotal,pensii,alocatii,ajutoare_stat,spatii_inchiriate,nr_persoane,Venit_per_membru,nr_credite

def filter_text_for_anexa2(text):
    start_index1 = text.find("Subsemnatul(a)") + len("Subsemnatul(a)")
    end_index1 = text.find(",student(a)")
    if start_index1 != -1 and end_index1 != -1:
        nume_student = text[start_index1:end_index1].strip()
    
    start_index2 = text.find("Facultății de") + len("Facultății de")
    end_index2 = text.find(",domeniul/specializarea")
    if start_index2 != -1 and end_index2 != -1:
        facultate = text[start_index2:end_index2].strip()

    start_index3 = text.find(",domeniul/specializarea") + len(",domeniul/specializarea")
    end_index3 = text.find(",în")
    if start_index3 != -1 and end_index3 != -1:
        specializare = text[start_index3:end_index3].strip()

    start_index4 = text.find("anul") + len("anul")
    end_index4 = text.find(",grupa")
    if start_index4 != -1 and end_index4 != -1:
        an = text[start_index4:end_index4].strip()

    start_index5 = text.find(",grupa") + len(",grupa")
    end_index5 = text.find(",media")
    if start_index5 != -1 and end_index5 != -1:
        grupa = text[start_index5:end_index5].strip()

    start_index6 = text.find(",media") + len(",media")
    end_index6 = text.find(",numărul")
    if start_index6 != -1 and end_index6 != -1:
        medie = text[start_index6:end_index6].strip()

    start_index7 = text.find(",CNP") + len(",CNP")
    end_index7 = text.find(",rog")
    if start_index7 != -1 and end_index7 != -1:
        CodNP_student = text[start_index7:end_index7].strip()

    start_index8 = text.find("universitar") + len("universitar")
    end_index8 = text.find(",a")
    if start_index8 != -1 and end_index8 != -1:
        AnUniv = text[start_index8:end_index8].strip()

    start_index9 = text.find("în perioada") + len("în perioada")
    end_index9 = text.find(",alte venituri")
    if start_index9 != -1 and end_index9 != -1:
        perioada = text[start_index9:end_index9].strip()

    start_index17 = text.find("de credite") + len("de credite")
    end_index17 = text.find(",CNP")
    if start_index17 != -1 and end_index17 != -1:
        nr_credite = text[start_index17:end_index17].strip()
    

    return nume_student, facultate, specializare, an, grupa, medie, CodNP_student, AnUniv,perioada,nr_credite

def filter_text_for_anexa3(text):
    start_index1 = text.find("Subsemnatul(a)") + len("Subsemnatul(a)")
    end_index1 = text.find(",student(a)")
    if start_index1 != -1 and end_index1 != -1:
        nume_student = text[start_index1:end_index1].strip()
    
    start_index2 = text.find("Facultății de") + len("Facultății de")
    end_index2 = text.find(",domeniul/specializarea")
    if start_index2 != -1 and end_index2 != -1:
        facultate = text[start_index2:end_index2].strip()

    start_index3 = text.find(",domeniul/specializarea") + len(",domeniul/specializarea")
    end_index3 = text.find(",în")
    if start_index3 != -1 and end_index3 != -1:
        specializare = text[start_index3:end_index3].strip()

    start_index4 = text.find("anul") + len("anul")
    end_index4 = text.find(",grupa")
    if start_index4 != -1 and end_index4 != -1:
        an = text[start_index4:end_index4].strip()

    start_index5 = text.find(",grupa") + len(",grupa")
    end_index5 = text.find(",media")
    if start_index5 != -1 and end_index5 != -1:
        grupa = text[start_index5:end_index5].strip()

    start_index6 = text.find(",media") + len(",media")
    end_index6 = text.find(",numărul")
    if start_index6 != -1 and end_index6 != -1:
        medie = text[start_index6:end_index6].strip()

    start_index7 = text.find(",CNP") + len(",CNP")
    end_index7 = text.find(",rog")
    if start_index7 != -1 and end_index7 != -1:
        CodNP_student = text[start_index7:end_index7].strip()

    start_index8 = text.find("universitar") + len("universitar")
    end_index8 = text.find(",a")
    if start_index8 != -1 and end_index8 != -1:
        AnUniv = text[start_index8:end_index8].strip()


    start_index9 = text.find(",a bursei") + len(",a bursei")
    end_index9 = text.find("Solicit")
    if start_index9 != -1 and end_index9 != -1:
        tip_bursa = text[start_index9:end_index9].strip()
    
    start_index10 = text.find("motive:") + len("motive:")
    end_index10 = text.find("Am luat la")
    if start_index10 != -1 and end_index10 != -1:
        motiv = text[start_index10:end_index10].strip()

    start_index17 = text.find("de credite") + len("de credite")
    end_index17 = text.find(",CNP")
    if start_index17 != -1 and end_index17 != -1:
        nr_credite = text[start_index17:end_index17].strip()


    return nume_student, facultate, specializare, an, grupa, medie, CodNP_student, AnUniv,tip_bursa,motiv,nr_credite

def filter_acord_membrii_fam(text):
    start_index1 = text.find("Subsemnatul/subsemnata") + len("Subsemnatul/subsemnata")
    end_index1 = text.find(",născut/născută")
    if start_index1 != -1 and end_index1 != -1:
        nume_parinte = text[start_index1:end_index1].strip()
    
    start_index2 = text.find("data de") + len("data de")
    end_index2 = text.find(",în localitatea")
    if start_index2 != -1 and end_index2 != -1:
        data_nastere = text[start_index2:end_index2].strip()

    start_index3 = text.find(",CNP ") + len(",CNP ")
    end_index3 = text.find(",în calitate")
    if start_index3 != -1 and end_index3 != -1:
        CodNP_parinte = text[start_index3:end_index3].strip()

    start_index4 = text.find("calitate de") + len("calitate de")
    end_index4 = text.find("al/a studentului/studentei")
    if start_index4 != -1 and end_index4 != -1:
        tip_membru = text[start_index4:end_index4].strip()

    start_index5 = text.find("studentului/studentei") + len(",studentului/studentei")
    end_index5 = text.find(",în cadrul")
    if start_index5 != -1 and end_index5 != -1:
        nume_student = text[start_index5:end_index5].strip()

    start_index6 = text.find("Facultatea") + len("Facultatea")
    end_index6 = text.find(",ciclul")
    if start_index6 != -1 and end_index6 != -1:
        facultate = text[start_index6:end_index6].strip()

    start_index7 = text.find("studii/specializarea") + len("studii/specializarea")
    end_index7 = text.find(",forma de învățământ")
    if start_index7 != -1 and end_index7 != -1:
        specializare = text[start_index7:end_index7].strip()

    start_index8 = text.find("bursei") + len("bursei")
    end_index8 = text.find("În temeiul")
    if start_index8 != -1 and end_index8 != -1:
        tip_bursa = text[start_index8:end_index8].strip()

    start_index9 = text.find("stabil/reședința în") + len("stabil/reședința în")
    end_index9 = text.find(",telefon")
    if start_index9 != -1 and end_index9 != -1:
        adresa = text[start_index9:end_index9].strip()

    start_index10 = text.find(",telefon") + len(",telefon")
    end_index10 = text.find(",posesor")
    if start_index10 != -1 and end_index10 != -1:
        telefon_parinte = text[start_index10:end_index10].strip()

    return nume_parinte,data_nastere,CodNP_parinte,tip_membru,nume_student,facultate,specializare,tip_bursa,adresa,telefon_parinte

def filter_acord_solicitant(text):
    start_index1 = text.find("Subsemnatul/subsemnata") + len("Subsemnatul/subsemnata")
    end_index1 = text.find(",student/studentă")
    if start_index1 != -1 and end_index1 != -1:
        nume_student = text[start_index1:end_index1].strip()

    start_index2 = text.find("Facultății") + len("Facultății")
    end_index2 = text.find(",ciclul de studii")
    if start_index2 != -1 and end_index2 != -1:
        facultate = text[start_index2:end_index2].strip()
    
    start_index3 = text.find("fiul/fiica lui") + len("fiul/fiica lui")
    end_index3 = text.find("și al/a")
    if start_index3 != -1 and end_index3 != -1:
        mama_student = text[start_index3:end_index3].strip()

    start_index4 = text.find("și al/a") + len("și al/a")
    end_index4 = text.find(",cu domiciliul")
    if start_index4 != -1 and end_index4 != -1:
        tata_student = text[start_index4:end_index4].strip()

    start_index5 = text.find("actual în") + len("actual în")
    end_index5 = text.find(",telefon")
    if start_index5 != -1 and end_index5 != -1:
        domiciliu_student = text[start_index5:end_index5].strip()

    start_index6 = text.find(",telefon") + len(",telefon")
    end_index6 = text.find("posesor")
    if start_index6 != -1 and end_index6 != -1:
        telefon_student = text[start_index6:end_index6].strip()
    

    return nume_student,facultate,mama_student,tata_student,domiciliu_student,telefon_student

def filter_text_anexa4(text):
    start_index1 = text.find("Subsemnatul,") + len("Subsemnatul,")
    end_index1 = text.find(",student")
    if start_index1 != -1 and end_index1 != -1:
        nume_student = text[start_index1:end_index1].strip()
    
    start_index2 = text.find("Facultății de") + len("Facultății de")
    end_index2 = text.find(",programul")
    if start_index2 != -1 and end_index2 != -1:
        facultate = text[start_index2:end_index2].strip()

    start_index3 = text.find("studii") + len("studii")
    end_index3 = text.find(",anul")
    if start_index3 != -1 and end_index3 != -1:
        program_studii = text[start_index3:end_index3].strip()

    start_index4 = text.find("universitar") + len("universitar")
    end_index4 = text.find("Doresc")
    if start_index4 != -1 and end_index4 != -1:
        AnUniv = text[start_index4:end_index4].strip()

    start_index5 = text.find("deschis la") + len("deschis la")
    end_index5 = text.find(",titular")
    if start_index5 != -1 and end_index5 != -1:
        cont_deschis_la = text[start_index5:end_index5].strip()

    start_index6 = text.find(",titular") + len(",titular")
    end_index6 = text.find(",IBAN")
    if start_index6 != -1 and end_index6 != -1:
        titular = text[start_index6:end_index6].strip()
    

    start_index7 = text.find(",IBAN") + len(",IBAN")
    end_index7 = text.find(",urmând")
    if start_index7 != -1 and end_index7 != -1:
        iban = text[start_index7:end_index7].strip()
    
    return nume_student,facultate,program_studii,AnUniv,cont_deschis_la,titular,iban


def extract_text_from_image(image):
    # Deschide imaginea utilizând PIL
    img = Image.open(image)

    # Extrage textul din imagine
    text = pytesseract.image_to_string(img)
    return text

def extract_text_from_docx(docx_file):
    # Deschide documentul Word utilizând python-docx
    doc = Document(docx_file)

    # Extrage textul din document și din fiecare tabel
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"

    return text

def extract_text_from_pdf(pdf_file):
    # Deschide PDF-ul utilizând PyPDF2
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    
    # Extrage textul din fiecare pagină
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()

    return text


def identify_document_type(text):
    if "Anexa nr. 1" in text:
        return "Anexa nr. 1 pentru bursa sociala"
    if "Anexa nr. 2" in text:
        return "Anexa nr. 2 pentru bursa sociala"
    if "Anexa nr. 3" in text:
        return "Anexa nr. 3 pentru toate tipurile de burse, cu excepția bursei sociale"
    if "caracter personal al membrilor familiei" in text:
        return "Acord de prelucrarea datelor cu caracter personal al membrilor familiei studentului solicitant de bursă socială"
    if "solicitantului" in text:
        return "Acord de prelucrarea datelor cu caracter personal al solicitantului de bursă"
    if "Anexa Nr. 4" in text:
        return "Declarație de acceptare a bursei de Performanță"
    else:
        return "Tip de document necunoscut"

if __name__ == '__main__':
    app.run(debug=True)